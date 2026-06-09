from docx import Document
import re

def extract_risk_substances(doc):
    """提取风险物质相关内容"""
    risk_info = {
        'sections': [],
        'risk_categories': [],
        'common_risk_substances': [],
        'assessment_methods': []
    }
    
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 查找章节标题
        if text.startswith('一、') or text.startswith('二、') or text.startswith('三、') or \
           text.startswith('四、') or text.startswith('五、') or text.startswith('1.') or \
           text.startswith('2.') or text.startswith('3.'):
            current_section = text
            risk_info['sections'].append(text)
            print(f"章节: {text}")
        
        # 查找风险物质类别
        if '风险物质' in text or '杂质' in text or '污染物' in text:
            if current_section:
                risk_info['risk_categories'].append((current_section, text))
                print(f"  相关内容: {text[:100]}...")
        
        # 查找具体的风险物质名称
        # 常见风险物质模式
        risk_patterns = [
            r'铅|砷|汞|镉|铬|镍|钴',  # 重金属
            r'甲醇|甲醛|乙醛',           # 挥发性有机物
            r'苯|苯酚|氯苯',             # 苯系物
            r'二𫫇烷|亚硝胺',           # 致癌物质
            r'塑化剂|邻苯二甲酸',        # 塑化剂
            r'微生物|细菌|霉菌',         # 微生物
        ]
        
        for pattern in risk_patterns:
            if re.search(pattern, text):
                risk_info['common_risk_substances'].append(text)
    
    return risk_info

# 加载指导原则文件
guideline_path = r'D:\工作文档\法律法规文件\法律法规文件\20240430附件2《化妆品风险物质识别与评估技术指导原则》.docx'

try:
    doc = Document(guideline_path)
    print(f"成功加载文件: {guideline_path}")
    print(f"文档段落数: {len(doc.paragraphs)}")
    print(f"文档表格数: {len(doc.tables)}")
    print("\n" + "="*60)
    
    # 提取风险物质信息
    risk_info = extract_risk_substances(doc)
    
    print("\n" + "="*60)
    print("提取的风险物质类别:")
    for cat in risk_info['risk_categories'][:10]:
        print(f"- {cat[0]}: {cat[1][:80]}")
    
    print("\n" + "="*60)
    print("查找表格:")
    for ti, table in enumerate(doc.tables):
        if len(table.rows) > 0 and len(table.columns) > 0:
            headers = [cell.text.strip()[:30] for cell in table.rows[0].cells]
            print(f"表格 {ti}: {len(table.rows)}行 x {len(table.columns)}列")
            print(f"  表头: {headers}")
            
            # 如果是风险物质相关表格，打印前几行
            if any('风险' in h or '物质' in h or '杂质' in h for h in headers):
                print("  内容示例:")
                for ri, row in enumerate(table.rows[1:min(5, len(table.rows))]):
                    cells = [cell.text.strip()[:30] for cell in row.cells]
                    print(f"    行{ri+1}: {cells}")
    
except Exception as e:
    print(f"读取文件失败: {e}")