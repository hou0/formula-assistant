from docx import Document
import re

def find_section(doc, section_title):
    """查找文档中的指定章节"""
    for i, para in enumerate(doc.paragraphs):
        if section_title in para.text:
            return i
    return -1

def extract_section_content(doc, start_idx, next_section=None):
    """提取章节内容"""
    content = []
    for i in range(start_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        # 如果遇到下一个章节标题，停止提取
        if next_section and next_section in text:
            break
        
        # 如果是标题格式（如"4.1"、"五、"等），停止提取
        if re.match(r'^\d+\.?\d*\s', text) or re.match(r'^[一二三四五六七八九十]+、', text):
            break
            
        if text:
            content.append((i, para))
    
    return content

def analyze_section(doc, section_title, next_section=None):
    """分析指定章节"""
    idx = find_section(doc, section_title)
    if idx == -1:
        print(f"未找到章节: {section_title}")
        return None
    
    print(f"\n{'='*60}")
    print(f"章节: {section_title}")
    print(f"位置: 第 {idx} 段")
    print(f"{'='*60}")
    
    content = extract_section_content(doc, idx, next_section)
    
    for i, para in content:
        pf = para.paragraph_format
        print(f"\n段落 {i}:")
        print(f"文本: {para.text[:200]}...")
        print(f"样式: {para.style.name}")
        print(f"对齐: {pf.alignment}")
        print(f"首行缩进: {pf.first_line_indent}")
        print(f"行距: {pf.line_spacing}")
        
        for j, run in enumerate(para.runs):
            f = run.font
            print(f"  Run {j}: 字体={f.name}, 字号={f.size}, 加粗={f.bold}, 斜体={f.italic}")
    
    return content

# 加载模板文件
template_path = r'C:\Users\houping\Nutstore\1\我的坚果云\化妆品安全评估\安全评估报告模版和安全评估基本结论模版\xxxx精华液完整版安评报告模版.docx'

try:
    doc = Document(template_path)
    
    # 打印所有段落标题
    print("文档结构概览:")
    print("="*60)
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and (text.startswith('一、') or text.startswith('二、') or text.startswith('三、') or 
                    text.startswith('四、') or text.startswith('五、') or text.startswith('六、') or
                    text.startswith('七、') or text.startswith('八、') or text.startswith('九、') or
                    text.startswith('十、') or text.startswith('十一、') or text.startswith('十二、') or
                    re.match(r'^\d+\.', text)):
            print(f"段落 {i}: {text}")
    
    print("\n" + "="*60)
    print("查找表格:")
    print("="*60)
    for ti, table in enumerate(doc.tables):
        print(f"表格 {ti}: {len(table.rows)}行 x {len(table.columns)}列")
        # 打印表格第一行作为表头
        if len(table.rows) > 0 and len(table.columns) > 0:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            print(f"  表头: {headers}")
    
    # 分析"四、配方中各成分的安全评估"部分
    analyze_section(doc, "四、配方中各成分的安全评估", "五、")
    
    # 分析"十一、参考文献"部分（如果存在）
    idx = find_section(doc, "十一、参考文献")
    if idx == -1:
        idx = find_section(doc, "参考文献")
    if idx != -1:
        analyze_section(doc, doc.paragraphs[idx].text.strip(), "十二、")
    
    # 查找参考文献引用格式示例
    print("\n" + "="*60)
    print("查找文中的参考文献引用格式:")
    print("="*60)
    
    ref_pattern = re.compile(r'\[(\d+)\]')
    for i, para in enumerate(doc.paragraphs):
        if ref_pattern.search(para.text):
            print(f"段落 {i}: {para.text[:200]}...")
    
    # 查看表格内容示例
    print("\n" + "="*60)
    print("表格内容示例（表3）:")
    print("="*60)
    for ti, table in enumerate(doc.tables):
        if len(table.rows) > 0 and len(table.columns) > 0:
            first_cell = table.rows[0].cells[0].text.strip()
            if '成分' in first_cell or '安全评估' in first_cell:
                print(f"表格 {ti} 内容:")
                for ri, row in enumerate(table.rows[:5]):  # 只显示前5行
                    cells = [cell.text.strip()[:30] for cell in row.cells]
                    print(f"  行{ri}: {cells}")
    
except Exception as e:
    print(f"读取文件失败: {e}")