import os
from docx import Document

def inspect_docx(file_path):
    """检查docx文件的字体设置"""
    doc = Document(file_path)
    
    print(f"=== 文档段落字体信息 ===")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        pf = para.paragraph_format
        print(f'P{i} style="{para.style.name}" align={pf.alignment} space_before={pf.space_before} space_after={pf.space_after}')
        for j, run in enumerate(para.runs):
            f = run.font
            sz = f.size
            name = f.name
            print(f'  R{j}: "{run.text[:80]}" bold={f.bold} italic={f.italic} sz={sz} name="{name}"')

    print(f"\n=== 文档表格字体信息 ===")
    for ti, table in enumerate(doc.tables):
        print(f'\n=== Table {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ===')
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, para in enumerate(cell.paragraphs):
                    txt = para.text.strip()
                    if not txt:
                        continue
                    pf = para.paragraph_format
                    print(f'  R{ri}C{ci}P{pi}: "{txt[:60]}" align={pf.alignment}')
                    for rj, run in enumerate(para.runs):
                        f = run.font
                        print(f'    Run{rj}: bold={f.bold} sz={f.size} name="{f.name}"')

# 检查文件格式
template_path = r'C:\Users\houping\Downloads\wKg8v2f8dQKAGPuqAAFqAP9DK-o368.doc'

if os.path.exists(template_path):
    if template_path.lower().endswith('.docx'):
        inspect_docx(template_path)
    elif template_path.lower().endswith('.doc'):
        print(f"警告: 文件 {template_path} 是 .doc 格式（旧版Word格式）")
        print("python-docx 库只支持 .docx 格式")
        print("请先用Microsoft Word打开此文件，然后另存为 .docx 格式")
        
        # 尝试查找docx版本
        docx_path = template_path[:-4] + '.docx'
        if os.path.exists(docx_path):
            print(f"\n找到对应的 .docx 文件: {docx_path}")
            print("正在分析...")
            inspect_docx(docx_path)
        else:
            print(f"\n未找到对应的 .docx 文件")
            print("请转换文件格式后再试")
else:
    print(f"错误: 文件不存在 - {template_path}")
