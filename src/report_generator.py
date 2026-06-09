"""
安全评估报告生成引擎 (PRD 5.9 重写)
=========================================
使用 python-docx 生成 Word 格式的化妆品安全评估报告（完整版）
匹配 Nutstore 模板：封面 → 目录 → 摘要 → 产品信息 → 产品配方(表1/表2)
→ 各成分安全评估(表3) → 危害识别(表4) → 风险控制 → 结论
→ 评估人员 → 参考资料 → 附件

新增功能：
1. 样式定制系统 - 支持统一配置字体、字号、颜色
2. 优化目录生成 - 使用简洁的TOC字段生成
3. 模板化段落和表格样式
4. 支持多种报告格式输出
"""

import io
import os
import json
from datetime import datetime
from typing import Optional, Dict, List, Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.shared import OxmlElement as SharedOxmlElement

import database as db
import safety_engine as se


# ── 样式配置系统 ──

class ReportStyle:
    """报告样式配置类，统一管理所有样式参数"""
    
    def __init__(self):
        # 字体配置
        self.font_ascii = 'Times New Roman'
        self.font_cjk = '宋体'
        self.font_heading = '方正小标宋简体'
        
        # 字号配置（单位：pt）
        self.size_title = 22      # 封面标题
        self.size_heading1 = 16   # 一级标题
        self.size_heading2 = 14   # 二级标题
        self.size_body = 16       # 正文（英文）
        self.size_body_cjk = 14   # 正文（中文）
        self.size_table = 12      # 表格内容
        self.size_table_header = 12  # 表格标题
        self.size_caption = 14    # 表格/图片说明
        
        # 颜色配置
        self.color_header_bg = 'D9E2F3'  # 表格表头背景色
        self.color_text = RGBColor(0, 0, 0)       # 正文颜色
        self.color_heading = RGBColor(0, 0, 0)    # 标题颜色
        
        # 段落配置
        self.line_spacing = 1.5       # 行距
        self.first_line_indent = 0.74 # 首行缩进（cm）
        self.space_after_body = 6     # 正文段后间距（pt）
        self.space_after_heading = 12 # 标题段后间距（pt）
        
        # 目录配置
        self.toc_depth = 3           # 目录深度
        self.toc_heading_size = 22   # 目录标题字号
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式，便于序列化"""
        return {
            'font_ascii': self.font_ascii,
            'font_cjk': self.font_cjk,
            'font_heading': self.font_heading,
            'size_title': self.size_title,
            'size_heading1': self.size_heading1,
            'size_heading2': self.size_heading2,
            'size_body': self.size_body,
            'size_body_cjk': self.size_body_cjk,
            'size_table': self.size_table,
            'size_table_header': self.size_table_header,
            'size_caption': self.size_caption,
            'color_header_bg': self.color_header_bg,
            'color_text': (self.color_text.rgbRed, self.color_text.rgbGreen, self.color_text.rgbBlue),
            'color_heading': (self.color_heading.rgbRed, self.color_heading.rgbGreen, self.color_heading.rgbBlue),
            'line_spacing': self.line_spacing,
            'first_line_indent': self.first_line_indent,
            'space_after_body': self.space_after_body,
            'space_after_heading': self.space_after_heading,
            'toc_depth': self.toc_depth,
            'toc_heading_size': self.toc_heading_size,
        }
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'ReportStyle':
        """从字典加载样式配置"""
        style = cls()
        style.font_ascii = data.get('font_ascii', 'Times New Roman')
        style.font_cjk = data.get('font_cjk', '宋体')
        style.font_heading = data.get('font_heading', '方正小标宋简体')
        style.size_title = data.get('size_title', 22)
        style.size_heading1 = data.get('size_heading1', 16)
        style.size_heading2 = data.get('size_heading2', 14)
        style.size_body = data.get('size_body', 16)
        style.size_body_cjk = data.get('size_body_cjk', 14)
        style.size_table = data.get('size_table', 12)
        style.size_table_header = data.get('size_table_header', 12)
        style.size_caption = data.get('size_caption', 14)
        style.color_header_bg = data.get('color_header_bg', 'D9E2F3')
        style.line_spacing = data.get('line_spacing', 1.5)
        style.first_line_indent = data.get('first_line_indent', 0.74)
        style.space_after_body = data.get('space_after_body', 6)
        style.space_after_heading = data.get('space_after_heading', 12)
        style.toc_depth = data.get('toc_depth', 3)
        style.toc_heading_size = data.get('toc_heading_size', 22)
        return style


# 创建默认样式实例
DEFAULT_STYLE = ReportStyle()


# ── Helpers ──

def _set_run_font(run, size_pt=12, size_cjk_pt=None, font_ascii='Times New Roman',
                  font_east_asia='宋体', bold=False):
    """Set run font with separate ASCII/sz and CJK/szCs sizes (template match).

    In the Nutstore template:
      - Body text: sz=32(16pt), szCs=28(14pt), eastAsia=宋体
      - Headings:  sz=32(16pt), szCs=32(16pt), eastAsia=宋体
      - Table data: sz=24(12pt), eastAsia=宋体
      - Table captions: sz=28(14pt), bold, eastAsia=宋体
    """
    if size_cjk_pt is None:
        size_cjk_pt = size_pt
    run.font.size = Pt(size_pt)          # sets w:sz (ASCII/Latin)
    run.font.bold = bold
    run.font.name = font_ascii           # sets w:ascii + w:hAnsi

    rpr = run._element.get_or_add_rPr()

    # East-Asian font
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_east_asia)

    # w:szCs (CJK font size)
    szCs_el = rpr.find(qn('w:szCs'))
    if szCs_el is None:
        szCs_el = OxmlElement('w:szCs')
        sz_el = rpr.find(qn('w:sz'))
        if sz_el is not None:
            sz_el.addnext(szCs_el)
        else:
            rpr.append(szCs_el)
    szCs_el.set(qn('w:val'), str(int(size_cjk_pt * 2)))

    # w:bCs (CJK bold)
    bCs_el = rpr.find(qn('w:bCs'))
    if bold:
        if bCs_el is None:
            bCs_el = OxmlElement('w:bCs')
            rpr.append(bCs_el)
    else:
        if bCs_el is not None:
            rpr.remove(bCs_el)


def _set_cell(cell, text='', bold=False, size=12, size_cjk=None,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              vertical='center', font_ascii='Times New Roman',
              font_east_asia='宋体', color=None):
    """Set cell text with template-matching format (default 12pt 宋体)."""
    if size_cjk is None:
        size_cjk = size
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    if vertical == 'center':
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = p.add_run(str(text) if text is not None else '')
    _set_run_font(run, size_pt=size, size_cjk_pt=size_cjk,
                  font_ascii=font_ascii, font_east_asia=font_east_asia,
                  bold=bold)
    if color:
        run.font.color.rgb = color
    return run


def _set_cell_shading(cell, color_hex='D9E2F3'):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


# ── 目录生成功能 ──

def _add_table_of_contents(doc, style: ReportStyle = DEFAULT_STYLE) -> None:
    """
    添加文档目录（Table of Contents）
    
    使用Word的TOC字段自动生成目录，支持：
    - 自定义目录深度（默认3级）
    - 超链接跳转（点击目录项可跳转到对应章节）
    - 目录标题样式定制
    
    Args:
        doc: Document对象
        style: ReportStyle对象，用于定制目录样式
    """
    # 添加目录标题
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.paragraph_format.space_after = Pt(36)
    
    run = toc_heading.add_run('目录')
    _set_run_font(run, 
                  size_pt=style.toc_heading_size, 
                  size_cjk_pt=style.toc_heading_size,
                  font_ascii=style.font_ascii,
                  font_east_asia=style.font_heading,
                  bold=True)
    
    # 添加目录字段（使用python-docx的内置方法）
    toc_paragraph = doc.add_paragraph()
    toc_paragraph.paragraph_format.line_spacing = style.line_spacing
    
    # 创建TOC字段
    run = toc_paragraph.add_run()
    fld_char_begin = run._element
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    
    # 添加instrText（目录字段指令）
    instr_text = SharedOxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = f'TOC \\o "1-{style.toc_depth}" \\h \\z \\u'
    
    # 在当前段落后添加instrText
    toc_paragraph._element.addnext(instr_text)
    
    # 添加分隔符
    fld_char_separate = SharedOxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')
    instr_text.addnext(fld_char_separate)
    
    # 添加结束标记段落
    end_paragraph = doc.add_paragraph()
    end_run = end_paragraph.add_run()
    fld_char_end = end_run._element
    fld_char_end.set(qn('w:fldCharType'), 'end')
    
    # 添加分页符
    doc.add_page_break()


def _update_toc(doc) -> None:
    """
    更新目录（在文档生成完成后调用，确保目录页码正确）
    
    注意：python-docx生成的TOC字段需要在Word中手动更新（按F9）
    或者使用COM接口自动更新（需要安装pywin32）
    """
    pass


# ── 样式定制辅助函数 ──

def _set_run_font_with_style(run, style: ReportStyle, text_type: str = 'body', bold: bool = False) -> None:
    """
    使用样式配置设置字体
    
    Args:
        run: Run对象
        style: ReportStyle对象
        text_type: 文本类型 ('body', 'heading1', 'heading2', 'table', 'caption', 'title')
        bold: 是否加粗
    """
    font_map = {
        'title': (style.font_heading, style.size_title, style.size_title),
        'heading1': (style.font_cjk, style.size_heading1, style.size_heading1),
        'heading2': (style.font_cjk, style.size_heading2, style.size_heading2),
        'body': (style.font_cjk, style.size_body, style.size_body_cjk),
        'table': (style.font_cjk, style.size_table, style.size_table),
        'caption': (style.font_cjk, style.size_caption, style.size_caption),
    }
    
    font_east_asia, size_pt, size_cjk_pt = font_map.get(text_type, font_map['body'])
    
    _set_run_font(run,
                  size_pt=size_pt,
                  size_cjk_pt=size_cjk_pt,
                  font_ascii=style.font_ascii,
                  font_east_asia=font_east_asia,
                  bold=bold)


def _para_with_style(doc, text: str = '', style: ReportStyle = DEFAULT_STYLE, 
                     text_type: str = 'body', bold: bool = False, 
                     align: int = WD_ALIGN_PARAGRAPH.LEFT, 
                     first_line_indent: Optional[float] = None) -> Any:
    """
    使用样式配置添加段落
    
    Args:
        doc: Document对象
        text: 段落文本
        style: ReportStyle对象
        text_type: 文本类型 ('body', 'heading1', 'heading2', 'table', 'caption', 'title')
        bold: 是否加粗
        align: 对齐方式
        first_line_indent: 首行缩进（cm），默认使用样式配置
    
    Returns:
        Paragraph对象
    """
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = style.line_spacing
    
    # 设置段后间距
    if text_type in ['heading1', 'heading2']:
        p.paragraph_format.space_after = Pt(style.space_after_heading)
    else:
        p.paragraph_format.space_after = Pt(style.space_after_body)
    
    # 设置首行缩进
    indent = first_line_indent if first_line_indent is not None else style.first_line_indent
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    
    # 添加文本
    run = p.add_run(text)
    _set_run_font_with_style(run, style, text_type, bold)
    
    return p


def _heading_with_style(doc, text: str, style: ReportStyle = DEFAULT_STYLE, 
                        level: int = 1) -> Any:
    """
    使用样式配置添加标题
    
    Args:
        doc: Document对象
        text: 标题文本
        style: ReportStyle对象
        level: 标题级别（1或2）
    
    Returns:
        Paragraph对象
    """
    text_type = 'heading1' if level == 1 else 'heading2'
    return _para_with_style(doc, text, style, text_type, bold=True, 
                            align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=0)


# ── 原有函数保留（兼容旧代码） ──

def _heading(doc, text, level=1):
    """Add heading matching template: Normal style, 16pt, bold."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    _set_run_font(run, size_pt=16, size_cjk_pt=16,
                  font_ascii='Times New Roman',
                  font_east_asia='宋体', bold=True)
    return p


def _para(doc, text='', bold=False, size=16, size_cjk=14, after=6,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          first_line_indent=None, font_east_asia='宋体'):
    """Add paragraph with template-matching format (sz=16pt, szCs=14pt, 宋体)."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    _set_run_font(run, size_pt=size, size_cjk_pt=size_cjk,
                  font_ascii='Times New Roman',
                  font_east_asia=font_east_asia, bold=bold)
    return p


def _add_table(doc, headers, rows, col_widths=None, caption_heading=None):
    """Add a formatted table with template-matching font sizes.

    - caption_heading: optional bold 14pt heading above table (e.g. '表1 产品配方表')
    - Headers: 12pt bold 宋体
    - Data: 12pt 宋体
    """
    if caption_heading:
        _para(doc, caption_heading, bold=True, size=14, size_cjk=14,
              after=3, font_east_asia='宋体')

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row — 12pt bold
    for ci, h in enumerate(headers):
        _set_cell(table.rows[0].cells[ci], h, bold=True, size=12,
                  font_east_asia='宋体')
        _set_cell_shading(table.rows[0].cells[ci], 'D9E2F3')

    # Data rows — 12pt normal
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            _set_cell(table.rows[ri + 1].cells[ci], val, size=12,
                      font_east_asia='宋体')

    # Column widths
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                if ci < len(row.cells):
                    row.cells[ci].width = Cm(w)

    return table


def _parse_composition(comp_json: str) -> list[dict]:
    """Parse composition JSON string into list of {name, percent}."""
    if not comp_json:
        return []
    try:
        return json.loads(comp_json)
    except (json.JSONDecodeError, TypeError):
        return []


def _compute_actual_components(formula_rows: list[dict]) -> list[dict]:
    """Break formula into actual individual components with concentrations.

    For raw materials with composition (e.g. 水70%+甘油30% at 2% addition):
      - 水: 2% * 0.70 = 1.40%
      - 甘油: 2% * 0.30 = 0.60%
    For pure materials (no composition): use added_percent directly.
    Returns list of {name, percent, source, inci}.
    """
    components = []
    for row in formula_rows:
        name = row.get('name_zh', '')
        added = row.get('added_percent', 0)
        comps = _parse_composition(row.get('composition', ''))
        if comps:
            for c in comps:
                actual = added * c.get('percent', 100) / 100.0
                components.append({
                    'name': c.get('name', ''),
                    'percent': actual,
                    'source': name,
                    'inci': c.get('inci', row.get('name_inci', '')),
                })
        else:
            components.append({
                'name': name,
                'percent': added,
                'source': name,
                'inci': row.get('name_inci', ''),
            })
    return components


def load_references_seed() -> dict:
    """Load references seed from userdata/references_seed.json."""
    try:
        seed_path = os.path.join(os.path.dirname(__file__), 'userdata', 'references_seed.json')
        with open(seed_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        refs = data.get('references', [])
        return {r['id']: r['citation'] for r in refs}
    except Exception:
        return {}


class RefRegistry:
    """Manage collected references and produce [N] markers.

    Standalone (does not depend on generate_safety_report closure).
    """

    def __init__(self, seed: dict | None = None):
        self.seed = seed or load_references_seed()
        self.collected: list[str] = []
        self.index_map: dict[str, int] = {}

    def add(self, ref_text: str) -> str:
        """Register a reference and return its [N] marker. Dedup by text."""
        if not ref_text:
            return ''
        ref_text = ref_text.strip()
        if ref_text in self.index_map:
            return f'[{self.index_map[ref_text]}]'
        for _cid, citation in self.seed.items():
            if ref_text in citation or citation in ref_text:
                if citation in self.index_map:
                    return f'[{self.index_map[citation]}]'
                idx = len(self.collected) + 1
                self.collected.append(citation)
                self.index_map[citation] = idx
                return f'[{idx}]'
        idx = len(self.collected) + 1
        self.collected.append(ref_text)
        self.index_map[ref_text] = idx
        return f'[{idx}]'


def _build_comp_text(r: dict, refs: RefRegistry) -> str:
    """Build assessment text for a single component result."""
    concentration = r.get('concentration', 0)
    banned = r.get('banned', {})
    restricted = r.get('restricted', {})
    usage = r.get('usage_reference', {})
    exposure = r.get('exposure', {})
    overall = r.get('overall_pass', False)
    toxicology = r.get('toxicology', {})

    parts = []

    ref_source = usage.get('source', '')
    if ref_source:
        if 'CIR' in ref_source or '美国化妆品' in ref_source:
            parts.append('美国化妆品原料评价委员会（CIR）评估结果显示')
            parts.append(refs.add('CIR'))
        elif 'SCCS' in ref_source:
            parts.append('欧洲消费者安全科学委员会（SCCS）评估结果显示')
            parts.append(refs.add('SCCS'))
        elif 'IFRA' in ref_source or '国际日用香料' in ref_source:
            parts.append('其使用符合国际日用香料香精协会（IFRA）要求')
            parts.append(refs.add('IFRA'))
        else:
            parts.append(f'参考{ref_source}')
            parts.append(refs.add(ref_source))

    if toxicology.get('noael') and toxicology.get('moes') is not None:
        parts.append(f'NOAEL={toxicology["noael"]} mg/kg bw/day，MOE={toxicology["moes"]:.2f}')
    elif toxicology.get('noael'):
        parts.append(f'NOAEL={toxicology["noael"]} mg/kg bw/day')
    if toxicology.get('loael'):
        parts.append(f'LOAEL={toxicology["loael"]} mg/kg bw/day')

    ref_max = usage.get('ref_max')
    if ref_max is not None:
        if restricted.get('restricted'):
            parts.append(f'在驻留类化妆品中的使用浓度不超过{ref_max}%')
        elif usage.get('within_range') is True or usage.get('within_range') is None:
            parts.append(f'在驻留类化妆品中的使用浓度为{ref_max}%')
    else:
        parts.append(f'该原料的添加量为{concentration}%')

    if restricted.get('restricted'):
        if restricted.get('concentration_ok') is True:
            parts.append(f'本配方{concentration}%在限用范围内，应用风险在可接受范围之内')
        elif restricted.get('concentration_ok') is False:
            parts.append(f'本配方{concentration}%超出限用要求')
    elif usage.get('within_range') is not False:
        parts.append('在本产品中应用风险在可接受范围之内')

    if banned.get('banned'):
        parts.append('该原料为禁用成分')

    sed_val = exposure.get('sed')
    if sed_val is not None:
        parts.append(f'SED={sed_val:.6f} mg/kg bw/day')

    ttc = r.get('ttc', {})
    if ttc and ttc.get('class') and not toxicology.get('noael'):
        cls = ttc.get('class')
        thresh = ttc.get('threshold')
        ttc_moe = ttc.get('moe')
        margin = ttc.get('margin')
        sed_ug = ttc.get('sed')
        if cls == 'exclusion':
            parts.append(f'Cramer 决策树分类属排除类（{ttc.get("reason")}），不适用TTC法')
        elif thresh and sed_ug and ttc_moe is not None:
            parts.append(
                f'Cramer 分类 Class {cls}（{ttc.get("source")}），'
                f'TTC={thresh}μg/(kg·d)，'
                f'SED={sed_ug:.4g}μg/(kg·d)，'
                f'MoE={ttc_moe:.1f}，{margin or "已评估"}'
            )
        else:
            parts.append(f'Cramer 分类 Class {cls}（{ttc.get("source")}），TTC={thresh}μg/(kg·d)')

    ra = r.get('read_across', {})
    if ra and ra.get('applicable') and not toxicology.get('noael'):
        ana = ra.get('analogue', '')
        noael_src = ra.get('noael_source', '')
        uf = ra.get('uncertainty_factor', 1)
        dq = ra.get('data_quality', '中')
        parts.append(
            f'采用 Read-across 类比评估：类比物「{ana}」'
            f'（{ra.get("justification", "")}），'
            f'NOAEL 数据源 {noael_src}，'
            f'不确定因子 {uf}（数据质量：{dq}）'
        )

    if overall and not banned.get('banned'):
        parts.append('在正常使用条件下，该成分的使用是安全的')
    elif not overall:
        parts.append('在正常使用条件下，该成分的安全性需要关注')

    return '，'.join(parts) + '。'


def build_section4_data(results: list[dict]) -> dict:
    """Build the assessment text for '四、配方中各成分的安全评估'.

    Returns: {
        'title': '四、配方中各成分的安全评估',
        'paragraphs': [
            {'type': 'single'|'composite_head'|'composite_sub',
             'material_idx': 1,
             'name': '...',           # for single / composite_sub
             'components': ['...'],   # for composite_head
             'text': '...',           # for single / composite_sub
            },
            ...
        ],
        'references': ['...', ...]   # collected refs in order
    }
    """
    from collections import OrderedDict
    grouped = OrderedDict()
    for r in results:
        fid = r.get('formula_id')
        if fid is None:
            fid = f'__flat_{id(r)}'
        if fid not in grouped:
            grouped[fid] = []
        grouped[fid].append(r)

    refs = RefRegistry()
    paragraphs = []
    material_idx = 0
    for _fid, group in grouped.items():
        material_idx += 1
        if len(group) == 1:
            r = group[0]
            ing_name = r.get('ingredient', '')
            text = _build_comp_text(r, refs)
            paragraphs.append({
                'type': 'single',
                'material_idx': material_idx,
                'name': ing_name,
                'text': text,
            })
        else:
            comp_names = [r.get('ingredient', '') for r in group]
            paragraphs.append({
                'type': 'composite_head',
                'material_idx': material_idx,
                'components': comp_names,
                'text': f'{"、".join(comp_names[:-1])}和{comp_names[-1]}的混合物。',
            })
            for r in group:
                comp_name = r.get('ingredient', '')
                text = _build_comp_text(r, refs)
                paragraphs.append({
                    'type': 'composite_sub',
                    'material_idx': material_idx,
                    'name': comp_name,
                    'text': text,
                })

    return {
        'title': '四、配方中各成分的安全评估',
        'paragraphs': paragraphs,
        'references': refs.collected,
    }


def _build_assessment_table_data(results: list[dict]) -> list[list]:
    """Build 表3 data rows from assessment results."""
    rows = []
    for i, r in enumerate(results):
        banned = r.get('banned', {})
        restricted = r.get('restricted', {})
        usage = r.get('usage_reference', {})
        exposure = r.get('exposure', {})

        # Build safety assessment text
        parts = []
        if banned.get('banned'):
            parts.append('该原料为禁用成分')
        if restricted.get('restricted'):
            if restricted.get('concentration_ok') is False:
                parts.append(f"限用浓度超标(限值{restricted.get('max_allowed', '?')}%)")
            else:
                parts.append(f"符合限用要求(≤{restricted.get('max_allowed', '?')}%)")
        if usage.get('within_range') is True:
            parts.append(f"使用量在参考范围内(参考{usage.get('ref_max', '?')}%)")
        elif usage.get('within_range') is False:
            parts.append(f"超出参考用量(参考上限{usage.get('ref_max', '?')}%)")
        else:
            parts.append('无量参考数据')

        sed = exposure.get('sed')
        if sed is not None:
            parts.append(f"SED={sed:.6f} mg/kg bw/day")

        overall = r.get('overall_pass', False)
        conclusion = '通过' if overall else '不通过'
        parts.append(f"结论: {conclusion}")

        # Reference source
        ref_src = usage.get('source', exposure.get('source', '—'))
        if not ref_src or ref_src == '—':
            ref_src = '—'

        rows.append([
            str(i + 1),
            r.get('ingredient', ''),
            f"{r.get('concentration', 0):.2f}%",
            '；'.join(parts),
            ref_src,
        ])
    return rows


def _build_hazard_table(results: list[dict], fmt: str) -> list[list]:
    """Build 表4 危害识别表 data rows.

    fmt controls level of detail:
      - '基础版 (仅标识)' → simple pass/fail per ingredient
      - '标准版 (含毒理学参考)' → include toxicology reference info
      - '完整版 (含文献引用)' → include literature citations
    """
    rows = []
    for i, r in enumerate(results):
        ingredient = r.get('ingredient', '')
        banned = r.get('banned', {})
        restricted = r.get('restricted', {})
        usage = r.get('usage_reference', {})
        exposure = r.get('exposure', {})
        overall = r.get('overall_pass', False)

        if fmt.startswith('基础版'):
            hazard = '通过安全评估' if overall else '存在风险'
            if banned.get('banned'):
                hazard = '禁用成分'
            elif restricted.get('restricted') and restricted.get('concentration_ok') is False:
                hazard = '限用超标'
            ref = '—'
        elif fmt.startswith('标准版'):
            parts = []
            toxicology = r.get('toxicology', {})
            if toxicology.get('noael'):
                parts.append(f"NOAEL={toxicology['noael']} mg/kg bw/day")
            if toxicology.get('moes'):
                parts.append(f"MOE={toxicology['moes']:.2f}")
            parts.append('通过' if overall else '不通过')
            hazard = '；'.join(parts) if parts else ('通过' if overall else '不通过')
            ref = usage.get('source', '—')
        else:  # 完整版
            parts = []
            toxicology = r.get('toxicology', {})
            if toxicology.get('noael'):
                parts.append(f"NOAEL={toxicology['noael']} mg/kg bw/day")
            if toxicology.get('loael'):
                parts.append(f"LOAEL={toxicology['loael']} mg/kg bw/day")
            if toxicology.get('moes'):
                parts.append(f"MOE={toxicology['moes']:.2f}")
            parts.append('通过' if overall else '不通过')
            hazard = '；'.join(parts) if parts else ('通过' if overall else '不通过')
            refs = []
            if usage.get('source'):
                refs.append(usage['source'])
            if exposure.get('source'):
                refs.append(exposure.get('source', ''))
            ref = '；'.join(refs) if refs else '—'

        rows.append([ingredient, hazard, ref])

    # If no rows, add a placeholder
    if not rows:
        rows.append(['—', '—', '—'])

    return rows


# ── Main generation ──

def generate_safety_report(
    product_name: str,
    product_category: str = '',
    application_site: str = '',
    assessment_result: dict | None = None,
    manufacturer: str = '',
    applicant: str = '',
    report_date: str = '',
    report_number: str = '',
    market: str = 'CN',
    extra_info: dict | None = None,
    formula: list[dict] | None = None,
) -> bytes:
    """Generate a Word document safety assessment report (完整版).

    Matches Nutstore template structure exactly with 11 sections:
    封面 → 一、摘要 → 二、产品信息 → 三、产品配方(表1/表2)
    → 四、各成分安全评估(评估正文) → 五、可能存在的风险物质(表3)
    → 六、风险控制措施及建议 → 七、安全评估结论
    → 八、安全评估人员及签名 → 九、参考文献 → 十、附件
    """
    # ── Defaults ──
    if assessment_result is None:
        assessment_result = {}
    if extra_info is None:
        extra_info = {}
    if formula is None:
        formula = db.get_current_formula() or []

    results = assessment_result.get('results', [])
    total = assessment_result.get('total_ingredients', 0)
    passed = assessment_result.get('passed', 0)
    failed = assessment_result.get('failed', 0)
    all_pass = assessment_result.get('all_pass', False)

    usage_method = extra_info.get('usage_method', '')
    warning_label = extra_info.get('warning_label', '')
    daily_amount = extra_info.get('daily_amount', 0) or 0
    retention_factor = extra_info.get('retention_factor', 1.0)

    if not report_date:
        report_date = datetime.now().strftime('%Y年%m月%d日')
    if not report_number:
        report_number = f'SA-{datetime.now().strftime("%Y%m%d")}-001'

    # ══════════════════════════════════════════
    #  v2 S1: 参考文献 [N] 编号系统
    #  数据源: userdata/references_seed.json
    # ══════════════════════════════════════════
    _ref_registry = RefRegistry()
    add_ref = _ref_registry.add
    collected_refs = _ref_registry.collected

    # ══════════════════════════════════════════
    #  v2 S2: 风险控制措施动态生成
    #  数据源: userdata/risk_control_rules.json
    # ══════════════════════════════════════════
    def _load_risk_control_rules():
        try:
            rules_path = os.path.join(os.path.dirname(__file__), 'userdata', 'risk_control_rules.json')
            with open(rules_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return {'categories': [], 'default_measures': []}

    _RISK_RULES = _load_risk_control_rules()

    def _build_control_measures(product_class=None, product_category=''):
        """根据产品分类/产品类别/extra_info 触发条件, 动态生成六、风险控制措施。
        Returns: list of (paragraph_text, is_subheader) tuples.
        """
        ef = extra_info or {}
        if product_class is None:
            product_class = ef.get('product_class', '')
        out = []
        triggered = []

        for cat in _RISK_RULES.get('categories', []):
            trig = cat.get('trigger_match', {})
            extras = trig.get('extra_info', [])
            fallback = cat.get('trigger_fallback', {})
            fallback_cats = fallback.get('category_contains', [])

            matched = False
            for k in extras:
                if ef.get(k):
                    matched = True
                    break
            if not matched and fallback_cats:
                for kw in fallback_cats:
                    if kw in product_category:
                        matched = True
                        break

            if matched:
                triggered.append(cat)

        if triggered:
            for cat in triggered:
                label = cat.get('label', '')
                out.append((f'【{label}】', True))
                for m in cat.get('measures', []):
                    out.append((m, False))
                ref = cat.get('ref', '')
                if ref:
                    ref_marker = add_ref(ref)
                    if ref_marker:
                        out.append((f'相关法规及文献：{ref_marker}', False))
            out.append(('【通用风险控制措施】', True))
            for m in _RISK_RULES.get('default_measures', []):
                out.append((m, False))
        else:
            for m in _RISK_RULES.get('default_measures', []):
                out.append((m, False))

        return out

    # Compute SED for first result as reference
    retention_display = retention_factor
    for r in results:
        exp = r.get('exposure', {})
        if exp.get('sed') is not None:
            if daily_amount > 0 or retention_factor != 1.0:
                retention_display = exp.get('retention_factor', retention_factor)
            break

    # ══════════════════════════════════════════
    #  Document Setup
    # ══════════════════════════════════════════
    doc = Document()

    # Section margins - matching Nutstore template (A4, ~2.7cm sides, ~3.4cm top, ~3.2cm bottom)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.4)
    section.bottom_margin = Cm(3.2)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)

    # Default style
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    # ══════════════════════════════════════════
    #  COVER PAGE
    # ══════════════════════════════════════════
    _para(doc, '', after=0)
    _para(doc, '化妆品安全评估报告（完整版）', bold=True, size=22,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
    # Read cover fields from extra_info with fallbacks
    ef = extra_info or {}
    cover_applicant = ef.get('applicant', '') or applicant or manufacturer or '—'
    cover_addr = ef.get('applicant_addr', '') or '—'
    cover_manufacturer = ef.get('manufacturer', '') or manufacturer or '—'
    cover_evaluator = ef.get('evaluator', '') or ''
    cover_qs = ef.get('qs_director', '') or ''
    cover_report_date = ef.get('report_date', '') or report_date
    cover_selfcheck = ef.get('selfcheck_date', '') or ''

    cover_items = [
        ('产品名称：', product_name or '—'),
        ('备案人名称：', cover_applicant),
        ('备案人地址：', cover_addr),
        ('评估单位：', cover_manufacturer),
        ('评估人：', cover_evaluator),
        ('质量安全负责人：', cover_qs),
        ('评估日期：', cover_report_date),
        ('自查日期：', cover_selfcheck),
    ]
    for label, val in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        run_label = p.add_run(label)
        run_label.font.size = Pt(16)
        run_label.font.name = 'Times New Roman'
        run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_val = p.add_run(val)
        run_val.font.size = Pt(16)
        run_val.font.name = 'Times New Roman'
        run_val._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ══════════════════════════════════════════
    #  TABLE OF CONTENTS (使用优化的目录生成函数)
    # ══════════════════════════════════════════
    _add_table_of_contents(doc)

    # ══════════════════════════════════════════
    #  一、摘要
    # ══════════════════════════════════════════
    _heading(doc, '一、摘要', level=1)

    site_desc = application_site or '相应部位'
    product_type = '驻留类' if retention_display >= 0.5 else '淋洗类'

    # Collect unique raw material names from formula
    material_names = []
    seen_names = set()
    for item in formula:
        name = (item.get('name_zh', '') or '').strip()
        if name and name not in seen_names:
            seen_names.add(name)
            material_names.append(name)
    total_materials = len(material_names)

    # Detect known risk substances from formula ingredients
    risk_set = set()
    for item in formula:
        name = (item.get('name_zh', '') or '')
        # 聚醚类/聚乙二醇/PEG类 → 可能含二𫫇烷
        if any(k in name for k in ['聚醚', '聚乙二醇', '泊洛沙姆']):
            risk_set.add('二𫫇烷')
        elif 'PEG' in name or 'peg' in name:
            risk_set.add('二𫫇烷')
        # 乙氧基化类 → 可能含二𫫇烷
        if 'ethoxylated' in name.lower() or '乙氧基化' in name:
            risk_set.add('二𫫇烷')
        # 甘油/二醇类 → 可能含二甘醇
        if any(k in name for k in ['甘油', '二醇']):
            risk_set.add('二甘醇')
        if 'glycol' in name.lower() and 'hexylene' not in name.lower() and 'caprylyl' not in name.lower():
            risk_set.add('二甘醇')
        # 乙醇/甲醇 → 甲醇
        if '甲醇' in name:
            risk_set.add('甲醇')
        if '乙醇' in name and '苯氧乙醇' not in name and '乙醇胺' not in name:
            risk_set.add('甲醇')
        # 香精 → 邻苯二甲酸酯
        if '香精' in name or 'parfum' in name.lower() or 'fragrance' in name.lower():
            risk_set.add('邻苯二甲酸酯')
    # 重金属总是需要评估
    risk_set.add('重金属（铅、汞、砷、镉）')
    risk_list = sorted(risk_set)
    risk_count = len(risk_list)

    # Build summary text matching Nutstore template format
    name_list_str = '、'.join(material_names)

    usage_desc = usage_method if usage_method else f'适用于{site_desc}，可每日使用'
    usage_desc = usage_desc.rstrip('。') if usage_desc else f'适用于{site_desc}，可每日使用'
    summary_text = (
        f'{product_name}为{product_type}化妆品，'
        f'{usage_desc}。'
        f'依据《化妆品安全评估技术导则》（2021年版）有关规定，'
        f'对产品的微生物、有害物质和稳定性等进行了检测，'
        f'并对配方所用的{name_list_str}等{total_materials}种成分进行评估，'
        f'可能存在的{"、".join(risk_list)}等{risk_count}种风险物质开展了安全评估。'
        f'结果显示，该产品在正常、合理及可预见的使用情况下，'
        f'不会对人体健康产生危害。'
    )
    _para(doc, summary_text, after=6, first_line_indent=0.74)

    # ══════════════════════════════════════════
    #  二、产品简介
    # ══════════════════════════════════════════
    _heading(doc, '二、产品简介', level=1)

    info_lines = [
        ('1', '产品名称', product_name or '—'),
        ('2', '使用方法', usage_method or '—'),
        ('3', '日均使用量（g/day）', f'{daily_amount}' if daily_amount > 0 else '*'),
        ('4', '驻留因子', str(retention_display)),
        ('5', '全身暴露量（SED）', 'SED=日均使用量×驻留因子×成分在配方中百分比×经皮吸收率÷体重'),
    ]
    for num, label, val in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f'{num}、{label}：{val}')
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    footnotes = []
    if daily_amount <= 0:
        footnotes.append(
            '* 日均使用量参考《THE SCCS NOTES OF GUIDANCE FOR THE TESTING OF '
            'COSMETIC INGREDIENTS AND THEIR SAFETY EVALUATION（12TH REVISION）》。'
        )
    footnotes.append('# 体重一般为默认的成人体重（60 kg）；经皮吸收率以100%计。')
    for fn in footnotes:
        _para(doc, fn, after=2, first_line_indent=0)

    # ══════════════════════════════════════════
    #  三、产品配方
    # ══════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, '三、产品配方', level=1)

    _para(doc,
          '本配方中所使用的原料均已列入《已使用化妆品原料目录》'
          '或《化妆品安全技术规范》（2015年版）。产品配方表见表1，'
          '产品实际成分含量表见表2。',
          after=6, first_line_indent=0.74)

    # 表1: 产品配方表
    _para(doc, '表1  产品配方表', bold=True, size=14, size_cjk=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

    table1_headers = ['序号', '中文名称', 'INCI名称/英文名称', '使用目的',
                      '在《已使用原料目录》中的序号', '备注']
    table1_rows = []
    # Group ingredients by name for combined rows matching template style
    seen_ingredients = {}
    for i, r in enumerate(results):
        name = r.get('ingredient', '')
        inci = r.get('inci_name', '')
        purpose = r.get('purpose', '—')
        conc = r.get('concentration', 0)
        code = r.get('used_raw_code', '')
        if name in seen_ingredients:
            seen_ingredients[name].append((inci, purpose, conc, code))
        else:
            seen_ingredients[name] = [(inci, purpose, conc, code)]

    idx = 1
    for ing_name, variants in seen_ingredients.items():
        if len(variants) == 1:
            inci, purpose, conc, code = variants[0]
            table1_rows.append([
                str(idx), ing_name, inci, purpose,
                code or '—', ''
            ])
            idx += 1
        else:
            first = True
            for inci, purpose, conc, code in variants:
                table1_rows.append([
                    str(idx) if first else '', ing_name, inci, purpose,
                    code or '—', ''
                ])
                first = False
            idx += 1

    if not table1_rows:
        table1_rows.append(['—', '—', '—', '—', '—', '—'])

    _add_table(doc, table1_headers, table1_rows,
               col_widths=[1.0, 3.0, 3.5, 2.0, 3.5, 2.0])

    _para(doc, '', after=6)

    # 表2: 实际成分含量
    _para(doc, '表2  实际成分含量', bold=True, size=14, size_cjk=14,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

    table2_headers = ['标准中文名称', 'INCI名', '实际成分含量（%）']
    components = _compute_actual_components(formula)
    # Deduplicate and sum percentages
    seen = {}
    for c in components:
        name = c['name']
        pct = c['percent']
        if name in seen:
            seen[name]['percent'] += pct
        else:
            seen[name] = {'percent': pct, 'inci': c.get('inci', c.get('source', ''))}
    table2_rows = []
    for name, data in sorted(seen.items(), key=lambda x: -x[1]['percent']):
        table2_rows.append([name, data['inci'], f'{data["percent"]:.2f}%'])
    if not table2_rows:
        table2_rows.append(['—', '—', '—'])

    _add_table(doc, table2_headers, table2_rows,
               col_widths=[4.0, 4.0, 2.5])

    # ══════════════════════════════════════════
    #  四、配方中各成分的安全评估
    # ══════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, '四、配方中各成分的安全评估', level=1)

    # Group results by formula_id — composite raw materials share the same id
    from collections import OrderedDict
    grouped = OrderedDict()
    for r in results:
        fid = r.get('formula_id')
        if fid is None:
            fid = f'__flat_{id(r)}'
        if fid not in grouped:
            grouped[fid] = []
        grouped[fid].append(r)


    def _write_para_run(p, text: str, bold: bool = False):
        """Add a formatted run to an existing paragraph."""
        r = p.add_run(text)
        r.font.size = Pt(16)
        r.font.bold = bold
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        return r

    material_idx = 0
    for fid, group in grouped.items():
        material_idx += 1

        if len(group) == 1:
            # ── Single-component: "N号原料：name，[assessment]" ──
            r = group[0]
            ing_name = r.get('ingredient', '')
            assessment_text = _build_comp_text(r, _ref_registry)

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Cm(0.74)
            _write_para_run(p, str(material_idx), bold=True)
            _write_para_run(p, '号原料：', bold=True)
            _write_para_run(p, f'{ing_name}，{assessment_text}')
        else:
            # ── Multi-component: raw material line + sub-paragraphs ──
            comp_names = [r.get('ingredient', '') for r in group]
            comp_display = '、'.join(comp_names[:-1]) + '和' + comp_names[-1]

            # "N号原料：comp1、comp2和comp3的混合物。"
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Cm(0.74)
            _write_para_run(p, str(material_idx), bold=True)
            _write_para_run(p, '号原料：', bold=True)
            _write_para_run(p, f'{comp_display}的混合物。')

            # Each component sub-paragraph: "comp_name，[assessment]"
            for r in group:
                comp_name = r.get('ingredient', '')
                assessment_text = _build_comp_text(r, _ref_registry)

                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.first_line_indent = Cm(0.74)
                _write_para_run(p, f'{comp_name}，')
                _write_para_run(p, assessment_text)

    # ══════════════════════════════════════════
    #  五、可能存在的风险物质的安全评估
    # ══════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, '五、可能存在的风险物质的安全评估', level=1)

    _para(doc,
          '本产品按照《化妆品安全评估技术导则》《化妆品风险物质识别'
          '与评估技术指导原则》的要求，基于当前科学认知水平，对可能由'
          '化妆品原料带入、生产过程中产生或带入的风险物质进行了评估，'
          '结果表明：',
          after=4, first_line_indent=0.74)
    _para(doc,
          '本产品的生产符合国家相关法律法规，对生产过程和产品包装材料'
          '进行严格的管理和控制。',
          after=4, first_line_indent=0.74)
    _para(doc,
          '产品中可能存在的安全性风险物质是技术上无法避免、由原料带入'
          '的杂质，残留的微量杂质在正常合理使用条件下不会对人体健康造成'
          '危害。产品安全性风险物质危害识别表见表3。',
          after=6, first_line_indent=0.74)

    # 表3: 安全性风险物质危害识别表
    _para(doc, '表3  安全性风险物质危害识别', bold=True, size=14, size_cjk=14,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

    hazard_headers = ['标准中文名称', '可能含有的风险物质', '备注']
    hazard_rows = []
    for r in results:
        ing = r.get('ingredient', '')
        risk_subs = r.get('risk_substances', [])

        if risk_subs:
            for rs in risk_subs:
                hazard_rows.append([rs, '见各成分安全评估章节', '符合限量要求'])
        else:
            hazard_rows.append([ing, '通过安全评估', '无风险'])

    if not hazard_rows:
        hazard_rows.append(['—', '—', '—'])

    _add_table(doc, hazard_headers, hazard_rows,
               col_widths=[3.0, 5.0, 3.0])

    _para(doc, '', after=6)
    _para(doc,
          '此外，该产品的检测报告显示，铅、砷、汞、镉、甲醇等有害物质'
          '含量符合《化妆品安全技术规范》（2015年版）第2章中有害物质'
          '限值要求。',
          after=4, first_line_indent=0.74)

    # ══════════════════════════════════════════
    #  六、风险控制措施或建议
    # ══════════════════════════════════════════
    _para(doc, '', after=6)
    _heading(doc, '六、风险控制措施或建议', level=1)

    site_desc = application_site or '面部'
    product_type2 = '驻留类' if retention_display >= 0.5 else '淋洗类'
    product_type_desc2 = f'{product_category or "面霜"}（{product_type2}化妆品）'
    _para(doc,
          f'本产品为{product_type_desc2}，适用于涂抹于{site_desc}，可每日使用。',
          after=4)

    if warning_label:
        _para(doc, f'本产品需标注警示用语：{warning_label}', after=4)
    else:
        _para(doc, '本产品无需标注警示用语。', after=4)

    if all_pass:
        _para(doc,
              '该产品配方中所有成分均符合《化妆品安全技术规范》（2015年版）'
              '及相关法规要求。在正常、合理及可预见的使用条件下，可以认为是安全的。',
              after=4, first_line_indent=0.74)
    else:
        _para(doc,
              f'该产品配方中共有 {failed} 种成分需要关注，建议如下：',
              after=4)
        for r in assessment_result.get('failed_ingredients', []):
            issues = '；'.join(r.get('issues', []))
            _para(doc, f'  · {r.get("ingredient", "")}：{issues}',
                  after=2)

    _para(doc, '', after=2)
    _para(doc, '风险控制措施：', bold=True, after=2)

    # v2 S2: 动态生成风险控制措施
    control_items = _build_control_measures(
        product_category=product_category,
    )
    for text, is_subheader in control_items:
        if is_subheader:
            _para(doc, text, bold=True, after=2)
        else:
            _para(doc, text, after=2)

    # ══════════════════════════════════════════
    #  七、安全评估结论
    # ══════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, '七、安全评估结论', level=1)

    _para(doc,
          f'本产品为{product_type_desc2}，可每日使用，涂抹于{site_desc}。'
          '主要暴露方式为经皮吸收，根据产品的特性，'
          '对本产品的暴露评估仅考虑经皮途径。',
          after=4, first_line_indent=0.74)

    # v2 S4: 气雾剂推进剂说明
    propellant_pct = (assessment_result or {}).get('propellant_percent', 0) or 0
    if propellant_pct > 0:
        _para(doc,
              f'本产品含推进剂（{propellant_pct:.1f}%），其他原料浓度已扣除推进剂后×100%计。',
              after=4, first_line_indent=0.74)

    _para(doc,
          '通过对产品以下各方面的综合评估：',
          after=6, first_line_indent=0.74)

    conclusion_items = [
        '各成分的安全评估结果显示，所有成分在本产品浓度下不会对人体健康产生危害；',
        '可能存在的安全性风险物质检测及评估结果显示，不会对人体健康产生危害；',
        '防腐剂挑战结果符合有关要求；',
        '微生物检验结果显示该产品微生物符合《化妆品安全技术规范》（2015年版）有关要求；',
        '有害物质检测结果显示，该产品有害物质含量符合《化妆品安全技术规范》（2015年版）有关要求；',
        '产品的理化特性、稳定性检测结果显示，符合相关要求；',
        '产品与包装材料的相容性评估结果显示，符合相关要求；',
        '配方中各成分之间未预见发生有害的相互作用。',
    ]
    for ci, item in enumerate(conclusion_items):
        _para(doc, f'{ci + 1}、{item}', after=2, first_line_indent=0)

    _para(doc, '', after=4)
    if all_pass:
        _para(doc,
              '综上，认为该产品在正常及合理、可预见的使用条件下，'
              '不会对人体健康产生危害。',
              after=6, first_line_indent=0.74, bold=True)
    else:
        _para(doc,
              f'综上，该产品配方中共有 {failed} 种成分存在安全性风险，'
              '建议对上述成分进行调整或替换，重新评估后方可上市。',
              after=6, first_line_indent=0.74, bold=True)

    # Assessment metrics table
    metric_rows = [
        ('配方原料总数', str(total)),
        ('通过安全评估', str(passed)),
        ('存在问题', str(failed)),
        ('评估人', ''),
        ('评估日期', report_date),
    ]
    _add_table(doc, ['项目', '内容'], metric_rows, col_widths=[4.5, 11.0])

    # ══════════════════════════════════════════
    #  八、安全评估人员签名
    # ══════════════════════════════════════════
    doc.add_page_break()
    _heading(doc, '八、安全评估人员签名', level=1)

    _para(doc, '评估人：     ', after=4)
    _para(doc, f'日期：{report_date}', after=4)
    _para(doc, '地址：     ', after=4)

    # ══════════════════════════════════════════
    _para(doc, '', after=6)
    _heading(doc, '九、安全评估人员简历', level=1)

    # Personnel info table matching template format
    person_table = doc.add_table(rows=6, cols=4)
    person_table.style = 'Table Grid'
    person_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    person_data = [
        ['姓名', extra_info.get('resume_name', ''), '性别', extra_info.get('resume_gender', '')],
        ['学位', extra_info.get('resume_degree', ''), '专业', extra_info.get('resume_major', '')],
        ['单位', extra_info.get('resume_org', ''), '', ''],
        ['个人简历', extra_info.get('resume_bio', ''), '', ''],
        ['从事工作简介', extra_info.get('resume_work', ''), '', ''],
        ['培训记录', extra_info.get('resume_training', ''), '', ''],
    ]
    for ri, row_data in enumerate(person_data):
        for ci, val in enumerate(row_data):
            _set_cell(person_table.rows[ri].cells[ci], val)
            if ci < 2:
                person_table.rows[ri].cells[ci].width = Cm(2.5)
    # Merge cells for columns 2-3 in rows 3-6
    for ri in range(2, 6):
        person_table.rows[ri].cells[1].merge(person_table.rows[ri].cells[3])

    # ══════════════════════════════════════════
    #  十、参考资料（v2 S1: 动态生成 [1]...[N]）
    _para(doc, '', after=6)
    _heading(doc, '十、参考资料', level=1)

    if collected_refs:
        for i, ref_text in enumerate(collected_refs, 1):
            _para(doc, f'[{i}] {ref_text}', after=2)
    else:
        _para(doc, '本次评估未引用外部参考文献。', after=2)

    _para(doc, '说明：参考文献按照格式要求列出明确的出处及来源，原文留档备查，无需提交。',
          after=4)

    # ══════════════════════════════════════════
    #  十一、附录
    doc.add_page_break()
    _heading(doc, '十一、附录', level=1)

    attachments = [
        '1、产品中二甘醇、苯酚检测报告',
        '2、防腐挑战测试评估结论',
        '3、稳定性测试评估结论',
        '4、包装材料相容性测试评估结论',
        '5、原料的工艺流程',
        '6、产品皮肤刺激性试验报告',
        '7、香精IFRA证书',
    ]
    for att in attachments:
        _para(doc, att, after=3)

    _para(doc, '', after=4)
    _para(doc,
          '注：此报告仅为示例，实际进行产品评估时，需按照《化妆品安全评估技术导则'
          '（2021年版）》结合产品的具体情况进行评估。',
          after=4)

    # ── Save to bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ── Utility functions ──

def build_3tables_data(formula: list[dict], results: list[dict]) -> dict:
    """Build data for the 3 main report tables.
    Returns: {
        'table1': {'title': '表1  产品配方表', 'headers': [...], 'rows': [[...]]},
        'table2': {'title': '表2  实际成分含量', ...},
        'table3': {'title': '表3  安全性风险物质危害识别', ...}
    }
    """
    # ── 表1: 产品配方表 ──
    seen_ingredients = {}
    for r in results:
        name = r.get('ingredient', '')
        inci = r.get('inci_name', '')
        purpose = r.get('purpose', '—')
        conc = r.get('concentration', 0)
        code = r.get('used_raw_code', '')
        if name in seen_ingredients:
            seen_ingredients[name].append((inci, purpose, conc, code))
        else:
            seen_ingredients[name] = [(inci, purpose, conc, code)]

    table1_rows = []
    idx = 1
    for ing_name, variants in seen_ingredients.items():
        if len(variants) == 1:
            inci, purpose, conc, code = variants[0]
            table1_rows.append([str(idx), ing_name, inci, purpose, code or '—', ''])
            idx += 1
        else:
            first = True
            for inci, purpose, conc, code in variants:
                table1_rows.append([str(idx) if first else '', ing_name, inci, purpose, code or '—', ''])
                first = False
            idx += 1
    if not table1_rows:
        table1_rows = [['—', '—', '—', '—', '—', '—']]

    # ── 表2: 实际成分含量 ──
    components = _compute_actual_components(formula)
    seen = {}
    for c in components:
        name = c['name']
        pct = c['percent']
        if name in seen:
            seen[name]['percent'] += pct
        else:
            seen[name] = {'percent': pct, 'inci': c.get('inci', c.get('source', ''))}
    table2_rows = []
    for name, data in sorted(seen.items(), key=lambda x: -x[1]['percent']):
        table2_rows.append([name, data['inci'], f'{data["percent"]:.2f}%'])
    if not table2_rows:
        table2_rows = [['—', '—', '—']]

    # ── 表 3: 安全性风险物质危害识别 ──
    # 加载风险物质规则
    risk_rules = {}
    try:
        risk_path = os.path.join(os.path.dirname(__file__), 'userdata', 'risk_substances.json')
        with open(risk_path, 'r', encoding='utf-8') as f:
            risk_data = json.load(f)
            # 建立风险物质名称到完整信息的映射
            for rs in risk_data.get('general_risk_substances', []):
                risk_rules[rs['name']] = rs
            for rs in risk_data.get('child_risk_substances', []):
                risk_rules[rs['name']] = rs
    except Exception:
        pass  # 如果加载失败，使用默认逻辑
    
    # 判断产品类型：牙膏还是化妆品
    # 通过产品名称或配方特征判断
    is_toothpaste = False
    if product_name and ('牙膏' in product_name or '口腔' in product_name):
        is_toothpaste = True
    
    # 使用表 2 的标准中文名称数据（来自_compute_actual_components）
    components = _compute_actual_components(formula)
    
    # 建立原料名称到标准中文名称的映射（用于查找风险物质）
    # 因为 risk_substances 是基于原料名称评估的
    raw_material_to_std = {}
    for comp in components:
        source = comp.get('source', '')  # 原料名称
        std_name = comp.get('name', '')  # 标准中文名称
        if source and std_name:
            if source not in raw_material_to_std:
                raw_material_to_std[source] = set()
            raw_material_to_std[source].add(std_name)
    
    # 收集所有需要显示的标准中文名称（去重）
    std_names_set = set()
    for comp in components:
        std_names_set.add(comp.get('name', ''))
    
    table3_rows = []
    # 遍历所有标准中文名称
    for std_name in sorted(std_names_set):
        # 查找该标准中文名称对应的原料（可能有多个原料对应同一个标准中文名称）
        matched_risk_substances = []
        
        # 反向查找：找到哪些原料包含这个标准中文名称
        for raw_name, std_names in raw_material_to_std.items():
            if std_name in std_names:
                # 在安全评估结果中查找该原料的风险物质
                for r in results:
                    if r.get('ingredient', '') == raw_name:
                        risk_subs = r.get('risk_substances', [])
                        if risk_subs:
                            matched_risk_substances.extend(risk_subs)
                        break
        
        # 去重
        matched_risk_substances = list(set(matched_risk_substances))
        
        if matched_risk_substances:
            # 有对应的风险物质，显示每个风险物质
            for risk_name in matched_risk_substances:
                # 从风险物质规则中获取详细信息
                risk_info = risk_rules.get(risk_name, {})
                full_description = risk_info.get('description', '符合限量要求')
                
                # 根据产品类型选择对应的描述
                # 描述格式："化妆品：xxx 描述。牙膏：xxx 描述。"
                if is_toothpaste:
                    # 牙膏产品，提取牙膏部分的描述
                    if '牙膏：' in full_description:
                        # 提取"牙膏："后面的内容
                        description = full_description.split('牙膏：')[1].split('。')[0] + '。'
                    else:
                        # 如果没有牙膏描述，使用完整描述
                        description = full_description
                else:
                    # 化妆品，提取化妆品部分的描述
                    if '化妆品：' in full_description:
                        # 提取"化妆品："后面的内容，到"牙膏："之前
                        description = full_description.split('化妆品：')[1].split('。牙膏：')[0] + '。'
                    else:
                        # 如果没有化妆品描述，使用完整描述
                        description = full_description
                
                table3_rows.append([std_name, risk_name, description])
        else:
            # 无对应的风险物质
            table3_rows.append([std_name, '无', '/'])

    if not table3_rows:
        table3_rows = [['—', '—', '—']]

    return {
        'table1': {
            'title': '表1  产品配方表',
            'headers': ['序号', '中文名称', 'INCI名称/英文名称', '使用目的', '在《已使用原料目录》中的序号', '备注'],
            'rows': table1_rows,
        },
        'table2': {
            'title': '表2  实际成分含量',
            'headers': ['标准中文名称', 'INCI名', '实际成分含量（%）'],
            'rows': table2_rows,
        },
        'table3': {
            'title': '表3  安全性风险物质危害识别',
            'headers': ['标准中文名称', '可能含有的风险物质', '备注'],
            'rows': table3_rows,
        },
    }


def save_report_to_file(filepath: str, report_bytes: bytes):
    """Save report bytes to file."""
    with open(filepath, 'wb') as f:
        f.write(report_bytes)


def generate_report_from_current(
    product_name: str,
    product_category: str = '',
    application_site: str = '',
    manufacturer: str = '',
    applicant: str = '',
    extra_info: dict | None = None,
) -> bytes:
    """Generate report from current formula and assessment.

    Convenience function: loads formula, runs assessment, generates report.
    """
    formula = db.get_current_formula()
    if not formula:
        raise ValueError('当前配方为空，无法生成报告')

    result = se.assess_formula(formula, product_category, application_site)

    return generate_safety_report(
        product_name=product_name,
        product_category=product_category,
        application_site=application_site,
        assessment_result=result,
        manufacturer=manufacturer,
        applicant=applicant,
        extra_info=extra_info,
        formula=formula,
    )


def format_assessment_as_text(result: dict) -> str:
    """Format assessment result as readable text (for preview)."""
    lines = []
    lines.append(f"安全评估结果: {result.get('passed', 0)}/{result.get('total_ingredients', 0)} 通过")
    if result.get('failed', 0) > 0:
        lines.append(f"存在问题: {result['failed']} 个原料")
    lines.append('')
    lines.append(f'{"原料名称":<20} {"INCI":<25} {"添加量":<8} {"禁用":<6} {"限用":<6} {"使用参考":<12} {"SED":<12} {"结论":<6}')
    lines.append('-' * 100)

    for r in result.get('results', []):
        name = r.get('ingredient', '')[:18]
        inci = r.get('inci_name', '')[:23]
        conc = f"{r.get('concentration', 0):.2f}%"
        banned = '禁用!' if r.get('banned', {}).get('banned') else '✓'
        restricted = '超标!' if r.get('restricted', {}).get('concentration_ok') is False else (
            '⚠限用' if r.get('restricted', {}).get('restricted') else '✓')
        usage_ref = r.get('usage_reference', {}).get('summary', '—')[:10]
        sed = f"{r.get('exposure', {}).get('sed', '—')}"
        result_text = '✓' if r.get('overall_pass') else '✗'
        lines.append(f'{name:<20} {inci:<25} {conc:<8} {banned:<6} {restricted:<6} {usage_ref:<12} {sed:<12} {result_text:<6}')

    lines.append('')
    if result.get('all_pass'):
        lines.append('结论: 全部通过 ✓')
    else:
        lines.append(f"结论: {result.get('failed', 0)} 个原料存在问题 ✗")
        for r in result.get('failed_ingredients', []):
            lines.append(f"  · {r.get('ingredient', '')}: {'; '.join(r.get('issues', []))}")

    return '\n'.join(lines)
